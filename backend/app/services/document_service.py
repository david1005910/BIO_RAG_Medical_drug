"""문서 파싱 서비스 - Docling을 사용한 PDF, DOCX, HTML 파일 처리

특정 디렉토리의 문서를 읽어서 텍스트를 추출하고
RAG 검색에 사용할 수 있도록 처리합니다.

Docling: IBM Research의 문서 파싱 라이브러리
- PDF, DOCX, PPTX, HTML 등 다양한 형식 지원
- 고급 레이아웃 분석 및 테이블 구조화
- 수식, 코드, 이미지 분류 지원
"""
import logging
import os
from pathlib import Path
from typing import List, Optional, Dict, Any
from dataclasses import dataclass

from app.core.config import settings

logger = logging.getLogger(__name__)

# Docling 사용 가능 여부 확인
_DOCLING_AVAILABLE = False
try:
    from docling.document_converter import DocumentConverter
    _DOCLING_AVAILABLE = True
    logger.info("✅ Docling 라이브러리 사용 가능")
except ImportError:
    logger.warning("⚠️ Docling 라이브러리 없음, 기본 파서 사용")


@dataclass
class ParsedDocument:
    """파싱된 문서 데이터"""
    filename: str
    filepath: str
    content: str
    file_type: str
    page_count: int = 1
    metadata: Dict[str, Any] = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentService:
    """문서 파싱 및 처리 서비스 (Docling 기반)

    지원 형식:
    - PDF (.pdf) - 고급 레이아웃 분석
    - DOCX (.docx) - 단락 및 테이블 추출
    - PPTX (.pptx) - 슬라이드 텍스트 추출
    - HTML (.html, .htm) - 웹 문서 파싱
    - 텍스트 (.txt, .md) - 일반 텍스트

    Docling 기능:
    - OCR 지원 (스캔 PDF)
    - 테이블 구조화
    - 수식 인식
    - Markdown 변환
    """

    def __init__(self, documents_dir: Optional[str] = None):
        self.documents_dir = documents_dir or settings.DOCUMENTS_DIR
        self._ensure_dir_exists()
        self._converter = None
        if _DOCLING_AVAILABLE:
            self._init_docling()

    def _init_docling(self):
        """Docling 컨버터 초기화"""
        try:
            self._converter = DocumentConverter()
            logger.info("✅ Docling DocumentConverter 초기화 완료")
        except Exception as e:
            logger.error(f"❌ Docling 초기화 실패: {e}")
            self._converter = None

    def _ensure_dir_exists(self):
        """문서 디렉토리가 존재하는지 확인"""
        Path(self.documents_dir).mkdir(parents=True, exist_ok=True)

    def list_documents(self) -> List[str]:
        """디렉토리의 모든 지원 문서 목록 반환"""
        supported_extensions = {'.pdf', '.docx', '.pptx', '.html', '.htm', '.txt', '.md'}
        documents = []

        for root, _, files in os.walk(self.documents_dir):
            for file in files:
                if Path(file).suffix.lower() in supported_extensions:
                    filepath = os.path.join(root, file)
                    documents.append(filepath)

        return documents

    async def parse_document(self, filepath: str) -> Optional[ParsedDocument]:
        """단일 문서 파싱 (Docling 우선 사용)

        Args:
            filepath: 문서 파일 경로

        Returns:
            ParsedDocument 또는 None
        """
        if not os.path.exists(filepath):
            logger.error(f"파일을 찾을 수 없습니다: {filepath}")
            return None

        file_ext = Path(filepath).suffix.lower()
        filename = Path(filepath).name

        try:
            # Docling으로 파싱 시도 (PDF, DOCX, PPTX, HTML 지원)
            if self._converter and file_ext in {'.pdf', '.docx', '.pptx', '.html', '.htm'}:
                content, page_count = self._parse_with_docling(filepath)
            # 텍스트 파일은 직접 읽기
            elif file_ext in {'.txt', '.md'}:
                content, page_count = self._parse_text(filepath)
            # Docling 없으면 기본 파서 사용
            elif file_ext == '.pdf':
                content, page_count = self._parse_pdf_fallback(filepath)
            elif file_ext == '.docx':
                content, page_count = self._parse_docx_fallback(filepath)
            elif file_ext in {'.html', '.htm'}:
                content, page_count = self._parse_html_fallback(filepath)
            else:
                logger.warning(f"지원하지 않는 파일 형식: {file_ext}")
                return None

            return ParsedDocument(
                filename=filename,
                filepath=filepath,
                content=content,
                file_type=file_ext[1:],  # 점 제거
                page_count=page_count,
                metadata={
                    "file_size": os.path.getsize(filepath),
                    "modified_time": os.path.getmtime(filepath),
                    "parser": "docling" if self._converter else "fallback",
                }
            )

        except Exception as e:
            logger.error(f"문서 파싱 실패 ({filepath}): {e}")
            return None

    def _parse_with_docling(self, filepath: str) -> tuple[str, int]:
        """Docling을 사용한 문서 파싱"""
        logger.info(f"📄 Docling으로 파싱 중: {filepath}")

        result = self._converter.convert(filepath)
        doc = result.document

        # Markdown으로 변환 (구조화된 텍스트)
        content = doc.export_to_markdown()

        # 페이지 수 추정 (PDF의 경우)
        page_count = len(doc.pages) if hasattr(doc, 'pages') and doc.pages else 1

        logger.info(f"✅ Docling 파싱 완료: {len(content)} chars, {page_count} pages")
        return content, page_count

    def _parse_pdf_fallback(self, filepath: str) -> tuple[str, int]:
        """PDF 파일 파싱 (Docling 없을 때)"""
        from pypdf import PdfReader

        reader = PdfReader(filepath)
        pages = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())

        content = "\n\n".join(pages)
        return content, len(reader.pages)

    def _parse_docx_fallback(self, filepath: str) -> tuple[str, int]:
        """DOCX 파일 파싱 (Docling 없을 때)"""
        from docx import Document

        doc = Document(filepath)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # 테이블 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        content = "\n\n".join(paragraphs)
        return content, 1

    def _parse_html_fallback(self, filepath: str) -> tuple[str, int]:
        """HTML 파일 파싱 (Docling 없을 때)"""
        from bs4 import BeautifulSoup

        with open(filepath, 'r', encoding='utf-8') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        # 스크립트와 스타일 태그 제거
        for script in soup(["script", "style"]):
            script.decompose()

        content = soup.get_text(separator='\n', strip=True)
        return content, 1

    def _parse_text(self, filepath: str) -> tuple[str, int]:
        """텍스트 파일 파싱"""
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        return content.strip(), 1

    async def parse_all_documents(self) -> List[ParsedDocument]:
        """디렉토리의 모든 문서 파싱

        Returns:
            파싱된 문서 리스트
        """
        documents = self.list_documents()
        parsed = []

        for filepath in documents:
            doc = await self.parse_document(filepath)
            if doc:
                parsed.append(doc)
                logger.info(f"✅ 문서 파싱 완료: {doc.filename} ({len(doc.content)} chars)")

        logger.info(f"📚 총 {len(parsed)}/{len(documents)}개 문서 파싱 완료")
        return parsed

    def chunk_document(
        self,
        document: ParsedDocument,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> List[Dict[str, Any]]:
        """문서를 청크로 분할

        Args:
            document: 파싱된 문서
            chunk_size: 청크 크기 (문자 수)
            overlap: 오버랩 크기

        Returns:
            청크 리스트
        """
        content = document.content
        chunks = []
        start = 0
        chunk_id = 0

        while start < len(content):
            end = start + chunk_size

            # 문장 경계에서 자르기 시도
            if end < len(content):
                # 마침표, 물음표, 느낌표 뒤에서 자르기
                for i in range(end, max(start + chunk_size // 2, 0), -1):
                    if content[i] in '.?!。':
                        end = i + 1
                        break

            chunk_text = content[start:end].strip()

            if chunk_text:
                chunks.append({
                    "chunk_id": f"{document.filename}_{chunk_id}",
                    "document_id": document.filename,
                    "content": chunk_text,
                    "metadata": {
                        **document.metadata,
                        "chunk_index": chunk_id,
                        "start_char": start,
                        "end_char": end,
                    }
                })
                chunk_id += 1

            start = end - overlap

        return chunks


# 싱글톤 인스턴스
_document_service: Optional[DocumentService] = None


def get_document_service() -> DocumentService:
    """문서 서비스 싱글톤 반환"""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service
